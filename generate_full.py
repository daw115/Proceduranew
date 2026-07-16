#!/usr/bin/env python3
"""Generate procedural docs WITH ALL screen images"""
import re, os, subprocess
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString, Tag

ROOT = Path(__file__).resolve().parent
INPUT = ROOT / "PROCEDURA_IDCC_TSO_v6 copy.html"
OUTDIR = ROOT / "output"
OUTDIR.mkdir(exist_ok=True)

html = INPUT.read_text(encoding="utf-8")
soup = BeautifulSoup(html, "html.parser")

PROCESSED = set()
IMG_COUNT = 0

def clean(text):
    return re.sub(r"\s+", " ", text).strip()

def get_children_text(el):
    parts = []
    for c in el.children:
        if isinstance(c, NavigableString): parts.append(str(c))
        elif isinstance(c, Tag): parts.append(get_children_text(c))
    return "".join(parts)

def extract_all_images(elem):
    """Recursively extract ALL images (including from tables), removing them."""
    imgs = []
    to_remove = []
    for child in list(elem.children):
        if isinstance(child, Tag) and child.name == "img":
            src, alt = child.get("src", ""), child.get("alt", "")
            if src and "tiles/" not in src:
                imgs.append((src, alt))
                to_remove.append(child)
        elif isinstance(child, Tag):
            imgs.extend(extract_all_images(child))
    for node in to_remove:
        node.decompose()
    return imgs

def table_to_md_safe(table):
    """Convert table to MD, skipping img elements."""
    rows = []
    for tr in table.find_all("tr"):
        cells = []
        for td in tr.find_all(["td", "th"]):
            td_copy = Tag(name=td.name, attrs=td.attrs)
            for c in td.children:
                if isinstance(c, NavigableString):
                    td_copy.append(str(c))
                elif isinstance(c, Tag) and c.name != "img":
                    td_copy.append(Tag(name=c.name, attrs=c.attrs))
                    for gc in c.children:
                        if isinstance(gc, NavigableString):
                            td_copy.append(str(gc))
            cells.append(clean(td_copy.get_text()))
        if cells: rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 2:
        header = rows[0]
        sep = "| " + " | ".join(["---"] * len(rows[0].split("|"))) + " |"
        return "\n".join([header, sep] + rows[1:])
    return "\n".join(rows)

def extract_text_from_tree(elem):
    texts = []
    for child in elem.children:
        if isinstance(child, NavigableString):
            t = clean(str(child))
            if t: texts.append(t)
        elif isinstance(child, Tag):
            tag = child.name
            if tag in ("p","div","li"):
                for gc in child.children:
                    if isinstance(gc, NavigableString):
                        t = clean(str(gc))
                        if t: texts.append(t)
            elif tag in ("strong", "b"):
                texts.append("**" + get_children_text(child).strip() + "**")
            elif tag in ("em", "i"):
                texts.append("*" + get_children_text(child).strip() + "*")
            elif tag == "span":
                if not child.get("class") and not child.get("style"):
                    texts.append(get_children_text(child).strip())
            elif tag == "a":
                text = clean(child.get_text())
                if text: texts.append(text)
            elif tag == "code":
                texts.append("`" + get_children_text(child).strip() + "`")
            elif tag == "pre":
                code = clean(child.get_text())
                texts.append("```\n" + code + "\n```")
            else:
                texts.append(get_children_text(child).strip())
    return texts

# Remove nav, scripts, styles, images
for bad in soup(["nav", "script", "style", "link"]):
    bad.decompose()

main_div = soup.find('main')
content_el = main_div if main_div else soup.body

sections = []
for elem in content_el.descendants:
    if not isinstance(elem, Tag) or elem.name not in ("h1","h2","h3","h4"):
        continue
    title = clean(elem.get_text())
    h_level = int(elem.name[1])
    sec = {"level": h_level, "title": title, "content": [], "images": []}

    # Collect ALL siblings
    sib = elem.next_sibling
    while sib:
        if isinstance(sib, Tag) and sib.name in ("h1","h2","h3","h4"):
            break
        if isinstance(sib, Tag):
            # Extract images from ENTIRE subtree (including nested tables)
            imgs = extract_all_images(sib)
            sec["images"].extend(imgs)
            # Now extract text
            if sib.name == "table":
                sec["content"].append(("table", table_to_md_safe(sib)))
            elif sib.name in ("p","div","ul","ol","li"):
                sec["content"].extend(extract_text_from_tree(sib))
            elif sib.name == "br":
                sec["content"].append("")
            else:
                sec["content"].extend(extract_text_from_tree(sib))
        elif isinstance(sib, NavigableString):
            t = clean(str(sib))
            if t: sec["content"].append(t)
        sib = sib.next_sibling
    sections.append(sec)

# ── Markdown ─────────────────────────────────────────────────────────
md_lines = ["# Procedura IDCC TSO v6\n", "*Instrukcja operacyjna dyspozytora PSE — proces IDCC*\n"]

for sec in sections:
    lvl = sec["level"]
    if lvl == 0: md_lines.append(f"# {sec['title']}\n")
    elif lvl == 1: md_lines.append(f"## {sec['title']}\n")
    elif lvl == 2: md_lines.append(f"### {sec['title']}\n")
    else: md_lines.append(f"{'#' * (lvl + 1)} {sec['title']}\n")

    for item in sec["content"]:
        if isinstance(item, tuple): md_lines.append(item[1] + "\n")
        elif item: md_lines.append(f"{item}\n")

    for src, alt in sec["images"]:
        if src not in PROCESSED:
            PROCESSED.add(src)
            IMG_COUNT += 1
            src_path = ROOT / src
            if src_path.exists():
                dst = OUTDIR / src.replace("/", "_")
                dst.parent.mkdir(parents=True, exist_ok=True)
                dst.write_bytes(src_path.read_bytes())
                md_lines.append(f"\n![{alt}]({src})\n")
    md_lines.append("\n")

md = "\n".join(md_lines)
OUTDIR.joinpath("PROCEDURA_IDCC_TSO_v6.md").write_text(md, encoding="utf-8")
print(f"[+] Markdown: {OUTDIR}/PROCEDURA_IDCC_TSO_v6.md ({len(md)} bytes)")

# ── DOCX ─────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for sec in sections:
        lvl = sec["level"]
        hl = 1 if lvl <= 1 else (2 if lvl == 2 else (3 if lvl == 3 else 4))
        doc.add_heading(sec["title"], level=hl)

        for item in sec["content"]:
            if isinstance(item, tuple):
                _, tbl_md = item
                tbl_lines = tbl_md.strip().split("\n")
                if len(tbl_lines) >= 2:
                    rows = []
                    for tl in tbl_lines:
                        cells = [c.strip() for c in tl.strip("| ").split("|")]
                        cells = [c for c in cells if c]
                        if cells: rows.append(cells)
                    if rows:
                        ncols = len(rows[0])
                        table = doc.add_table(rows=len(rows), cols=ncols)
                        table.style = "Light Shading Accent 1"
                        for ri, row in enumerate(rows):
                            for ci, cell in enumerate(row):
                                if ci < len(row):
                                    table.cell(ri, ci).text = cell
                    continue
            else:
                if not item.strip():
                    doc.add_paragraph()
                    continue
                text = item
                if "**" in text:
                    parts = text.split("**")
                    para = doc.add_paragraph()
                    i = 0
                    while i < len(parts):
                        run = para.add_run(parts[i])
                        if i + 1 < len(parts):
                            i += 2
                            run.bold = True
                            run2 = para.add_run(parts[i])
                        i += 1
                else:
                    doc.add_paragraph(text)

        for src, alt in sec["images"]:
            if src not in PROCESSED:
                PROCESSED.add(src)
                IMG_COUNT += 1
                src_path = ROOT / src
                if src_path.exists():
                    fname = src.replace("/", "_")
                    dst = OUTDIR / fname
                    dst.parent.mkdir(parents=True, exist_ok=True)
                    dst.write_bytes(src_path.read_bytes())
                    doc.add_picture(str(dst), width=Pt(600))
                    doc.add_paragraph(alt)

    doc.save(str(OUTDIR / "PROCEDURA_IDCC_TSO_v6.docx"))
    print(f"[+] DOCX: {OUTDIR}/PROCEDURA_IDCC_TSO_v6.docx")
except ImportError:
    print("[!] python-docx not installed — skipping DOCX")

# ── PDF via WeasyPrint ───────────────────────────────────────────────
try:
    from weasyprint import HTML
    md_file = OUTDIR / "PROCEDURA_IDCC_TSO_v6.md"
    pdf_file = OUTDIR / "PROCEDURA_IDCC_TSO_v6.pdf"
    result = subprocess.run(
        ["pandoc", str(md_file), "-t", "html5", "--standalone",
         "--metadata", "title=Procedura IDCC TSO v6",
         "--css", "-"],
        input="""
        body { font-family: Arial, sans-serif; font-size: 11pt; margin: 2.5cm; }
        h1 { font-size: 18pt; color: #1a1a1a; }
        h2 { font-size: 15pt; color: #333; }
        h3 { font-size: 13pt; color: #444; }
        h4 { font-size: 11pt; color: #555; }
        img { max-width: 100%; height: auto; page-break-inside: avoid; }
        table { border-collapse: collapse; width: 100%; font-size: 10pt; }
        td, th { border: 1px solid #ddd; padding: 4px; }
        th { background: #f5f5f5; }
        """,
        capture_output=True, text=True, timeout=60,
    )
    if result.returncode == 0:
        HTML(string=result.stdout).write_pdf(str(pdf_file))
        print(f"[+] PDF: {pdf_file}")
    else:
        print(f"  [!] Pandoc HTML conversion failed: {result.stderr[:200]}")
except Exception as e:
    print(f"[!] PDF generation failed: {e}")

print(f"\nImages embedded: {IMG_COUNT}")
print(f"Sections processed: {len(sections)}")
print("Done.")
