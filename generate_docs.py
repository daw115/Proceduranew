#!/usr/bin/env python3
"""Generate procedural docs from PROCEDURA_IDCC_TSO_v6 copy.html"""
import re, sys, os, subprocess
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString, Tag

INPUT = Path("/Volumes/SSD/CLAUDE_WORK/Procedura/PROCEDURA_IDCC_TSO_v6 copy.html")
OUTDIR = Path("/Volumes/SSD/CLAUDE_WORK/Procedura/output")
OUTDIR.mkdir(exist_ok=True)

html = INPUT.read_text(encoding="utf-8")
soup = BeautifulSoup(html, "html.parser")

def clean(text):
    return re.sub(r"\s+", " ", text).strip()

def get_children_text(el):
    parts = []
    for c in el.children:
        if isinstance(c, NavigableString):
            parts.append(str(c))
        elif isinstance(c, Tag):
            parts.append(get_children_text(c))
    return "".join(parts)

def table_to_md(table):
    rows = []
    for tr in table.find_all("tr"):
        cells = []
        for td in tr.find_all(["td", "th"]):
            cells.append(clean(td.get_text()))
        if cells:
            rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 2:
        header = rows[0]
        sep = "| " + " | ".join(["---"] * len(rows[0].split("|"))) + " |"
        return "\n".join([header, sep] + rows[1:])
    return "\n".join(rows)

# Remove nav, scripts, styles, images
for bad in soup(["nav", "script", "style", "link", "img"]):
    bad.decompose()

# Find the main content area (nested in wrap > main)
main_div = soup.find('main')
if main_div:
    content_el = main_div
else:
    content_el = soup.body

sections = []
for elem in content_el.descendants:
    if not isinstance(elem, Tag):
        continue
    if elem.name in ("h1", "h2", "h3", "h4"):
        title = clean(elem.get_text())
        h_level = int(elem.name[1])
        sec = {"level": h_level, "title": title, "content": []}
        sib = elem.next_sibling
        while sib:
            if isinstance(sib, Tag) and sib.name in ("h1","h2","h3","h4"):
                break
            if isinstance(sib, Tag) and sib.name == "table":
                sec["content"].append(("table", table_to_md(sib)))
            elif isinstance(sib, Tag) and sib.name in ("p","div","ul","ol","li"):
                for child in sib.children:
                    if isinstance(child, NavigableString):
                        t = clean(str(child))
                        if t:
                            sec["content"].append(t)
                    elif isinstance(child, Tag):
                        tag = child.name
                        if tag in ("p","div","li"):
                            for gc in child.children:
                                if isinstance(gc, NavigableString):
                                    t = clean(str(gc))
                                    if t:
                                        sec["content"].append(t)
                        elif tag in ("strong", "b"):
                            sec["content"].append("**" + get_children_text(child).strip() + "**")
                        elif tag in ("em", "i"):
                            sec["content"].append("*" + get_children_text(child).strip() + "*")
                        elif tag == "span":
                            if not child.get("class") and not child.get("style"):
                                sec["content"].append(get_children_text(child).strip())
                        elif tag == "a":
                            text = clean(child.get_text())
                            if text:
                                sec["content"].append(text)
                        elif tag == "code":
                            sec["content"].append("`" + get_children_text(child).strip() + "`")
                        elif tag == "pre":
                            code = clean(child.get_text())
                            sec["content"].append("```\n" + code + "\n```")
                        else:
                            sec["content"].append(get_children_text(child).strip())
            elif isinstance(sib, Tag) and sib.name == "br":
                sec["content"].append("")
            elif isinstance(sib, NavigableString):
                t = clean(str(sib))
                if t:
                    sec["content"].append(t)
            sib = sib.next_sibling
        sections.append(sec)

# ── Markdown ─────────────────────────────────────────────────────────
md_lines = []
md_lines.append("# Procedura IDCC TSO v6\n")
md_lines.append("*Instrukcja operacyjna dyspozytora PSE — proces IDCC*\n")

for sec in sections:
    lvl = sec["level"]
    if lvl == 0:
        md_lines.append(f"# {sec['title']}\n")
    elif lvl == 1:
        md_lines.append(f"## {sec['title']}\n")
    elif lvl == 2:
        md_lines.append(f"### {sec['title']}\n")
    else:
        md_lines.append(f"{'#' * (lvl + 1)} {sec['title']}\n")
    for item in sec["content"]:
        if isinstance(item, tuple):
            md_lines.append(item[1] + "\n")
        elif item:
            md_lines.append(f"{item}\n")
    md_lines.append("\n")

md = "\n".join(md_lines)
OUTDIR.joinpath("PROCEDURA_IDCC_TSO_v6.md").write_text(md, encoding="utf-8")
print(f"[+] Markdown: {OUTDIR}/PROCEDURA_IDCC_TSO_v6.md ({len(md)} bytes)")

# ── DOCX ─────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for sec in sections:
        lvl = sec["level"]
        heading_level = 1 if lvl <= 1 else (2 if lvl == 2 else (3 if lvl == 3 else 4))
        p = doc.add_heading(sec["title"], level=heading_level)

        for item in sec["content"]:
            if isinstance(item, tuple):
                _, tbl_md = item
                tbl_lines = tbl_md.strip().split("\n")
                if len(tbl_lines) >= 2:
                    rows = []
                    for tl in tbl_lines:
                        cells = [c.strip() for c in tl.strip("| ").split("|")]
                        cells = [c for c in cells if c]
                        if cells:
                            rows.append(cells)
                    if rows:
                        ncols = len(rows[0])
                        table = doc.add_table(rows=len(rows), cols=ncols)
                        table.style = "Light Shading Accent 1"
                        for ri, row in enumerate(rows):
                            for ci, cell in enumerate(row):
                                if ci < len(row):
                                    table.cell(ri, ci).text = cell
                    continue
            else:
                if not item.strip():
                    doc.add_paragraph()
                    continue
                text = item
                if "**" in text:
                    parts = text.split("**")
                    para = doc.add_paragraph()
                    i = 0
                    while i < len(parts):
                        run = para.add_run(parts[i])
                        if i + 1 < len(parts):
                            i += 2
                            run.bold = True
                            run2 = para.add_run(parts[i])
                        i += 1
                else:
                    doc.add_paragraph(text)

    doc.save(str(OUTDIR / "PROCEDURA_IDCC_TSO_v6.docx"))
    print(f"[+] DOCX: {OUTDIR}/PROCEDURA_IDCC_TSO_v6.docx")
except ImportError:
    print("[!] python-docx not installed — skipping DOCX")

# ── PDF via pandoc ───────────────────────────────────────────────────
try:
    md_file = OUTDIR / "PROCEDURA_IDCC_TSO_v6.md"
    pdf_file = OUTDIR / "PROCEDURA_IDCC_TSO_v6.pdf"
    result = subprocess.run(
        [
            "pandoc", str(md_file),
            "-o", str(pdf_file),
            "--pdf-engine=lualatex",
            "-V", "geometry:a4paper,top=2.5cm,bottom=2.5cm,left=2.5cm,right=2.5cm",
            "-V", "mainfont=Arial",
            "-V", "fontsize=11pt",
            "--toc",
            "-V", "toc-depth=3",
        ],
        capture_output=True, text=True, timeout=120,
    )
    if pdf_file.exists():
        print(f"[+] PDF: {pdf_file}")
    else:
        print(f"  [!] PDF engine issue: {result.stderr[:300]}")
except Exception as e:
    print(f"[!] PDF generation failed: {e}")

print("\nDone. Files in:", OUTDIR)
