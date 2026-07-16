#!/usr/bin/env python3
"""Konwersja PROCEDURA_IDCC_TSO_v5_2.html → DOCX z bocznym TOC.

Kroki:
  1. Czyszczenie HTML z elementów przeglądarkowych (editor, sidebar, cover, skrypty).
  2. Pandoc → DOCX z nagłówkami H1–H4 (okienko nawigacji Word = boczny panel).
  3. Post-processing python-docx: pole TOC, marginesy, ramki i dopasowanie tabel,
     skalowanie obrazów do szerokości strony.
"""
import re, subprocess
from pathlib import Path

ROOT = Path(__file__).resolve().parent
SRC = ROOT / 'PROCEDURA_IDCC_TSO_v5_2.html'
TMP = ROOT / '_v5_2_for_docx.html'
OUT = ROOT / 'PROCEDURA_IDCC_TSO_v5_2.docx'

html = SRC.read_text(encoding='utf-8')

# ── 1. Czyszczenie ─────────────────────────────────────────────────────────
def strip_block(src, start_re):
    m = re.search(start_re, src)
    if not m: return src
    i = m.start()
    depth = 0
    j = i
    for mm in re.finditer(r'<(/?)(div|nav)[^>]*>', src[i:], re.I):
        if mm.group(1) == '': depth += 1
        else:
            depth -= 1
            if depth == 0:
                j = i + mm.end()
                break
    return src[:i] + src[j:]

html = strip_block(html, r'<nav class="sidebar"')
html = strip_block(html, r'<div class="cover"')
html = strip_block(html, r'<div class="editor-bar"')
html = strip_block(html, r'<div id="search-panel"')
html = re.sub(r'<div class="topbar">.*?</div>', '', html, flags=re.S)
html = re.sub(r'<input\b[^>]*>', '', html, flags=re.I)
html = re.sub(r'<label\b[^>]*>.*?</label>', '', html, flags=re.S|re.I)
html = re.sub(r'<script\b[^>]*>.*?</script>', '', html, flags=re.S|re.I)

# Zrzuty ekranu: nadaj szerokość 620px (pandoc: px/96 → cale), żeby mieściły się na A4
html = re.sub(r'<img src="(img_ccm/ccm-[^"]+)"([^>]*?)/?>',
              r'<img src="\1" width="620"\2>', html)

# Usuń colgroup z px — pandoc lepiej radzi sobie z auto-szerokością w DOCX
html = re.sub(r'<colgroup>.*?</colgroup>', '', html, flags=re.S)

TMP.write_text(html, encoding='utf-8')

# ── 2. Pandoc ──────────────────────────────────────────────────────────────
cmd = [
    'pandoc', str(TMP),
    '-f', 'html',
    '-t', 'docx',
    '--toc', '--toc-depth=4',
    '-o', str(OUT),
    '--metadata', 'title=PROCEDURA IDCC TSO v5.2 – pełna instrukcja z katalogiem CCM',
]
res = subprocess.run(cmd, capture_output=True, text=True, cwd=str(ROOT))
if res.stderr: print('pandoc:', res.stderr[-300:])
print(f'DOCX: {OUT}  ({OUT.stat().st_size:,} B)')

# ── 3. Post-processing ─────────────────────────────────────────────────────
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_BREAK

doc = Document(str(OUT))

# 3a. Marginesy A4
for sec in doc.sections:
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(1.8)
    sec.top_margin = sec.bottom_margin = Cm(1.8)

USABLE = Cm(21.0 - 2*1.8)  # 17.4 cm

# 3b. Tabele: ramki, 100% szerokości, czcionka 9 pt, nagłówek z tłem
def set_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'A8B8C8')

def set_full_width(tbl):
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')   # 100%
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'autofit')

def shade_cell(cell, hexclr):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexclr)

n_tables = 0
for tbl in doc.tables:
    n_tables += 1
    set_borders(tbl)
    set_full_width(tbl)
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    r.font.size = Pt(9)
            if ri == 0:
                shade_cell(cell, '003B75')
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
print(f'Tabel sformatowanych: {n_tables}')

# 3c. Obrazy: zmieść w szerokości strony (zachowaj proporcje)
n_img = 0
for shape in doc.inline_shapes:
    if shape.width > USABLE:
        ratio = USABLE / shape.width
        shape.height = Emu(int(shape.height * ratio))
        shape.width  = Emu(int(USABLE))
        n_img += 1
print(f'Obrazów przeskalowanych: {n_img} / {len(doc.inline_shapes)}')

# 3d. Pole TOC Word + wskazówka nawigacji na początku
def add_word_toc(paragraph):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-4" \h \z \u')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = 'Spis treści — kliknij prawym przyciskiem → „Aktualizuj pole" lub naciśnij F9'
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)

hint = doc.paragraphs[0].insert_paragraph_before(
    'Wskazówka: w MS Word włącz Widok → Okienko nawigacji (lub Ctrl+F), aby uzyskać '
    'boczny panel ze spisem treści opartym na nagłówkach 1–4.'
)
for r in hint.runs:
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x00, 0x3B, 0x75)

toc_p = doc.paragraphs[1].insert_paragraph_before('')
add_word_toc(toc_p)
br = doc.paragraphs[2].insert_paragraph_before('')
br.add_run().add_break(WD_BREAK.PAGE)

doc.save(str(OUT))
print(f'Final: {OUT.stat().st_size:,} B')
TMP.unlink(missing_ok=True)
